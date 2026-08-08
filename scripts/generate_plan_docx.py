"""Generate the project plan DOCX from the Markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_SOURCE = ROOT / "docs" / "项目开发计划书.md"
DEFAULT_OUTPUT = ROOT / "docs" / "项目开发计划书.docx"

HEADING_COLOR = RGBColor(0x2E, 0x74, 0xB5)
HEADING_DARK = RGBColor(0x1F, 0x4D, 0x78)
HEADER_FILL = "F2F4F7"


def set_east_asia(style_or_run, name: str = "Microsoft YaHei") -> None:
    rpr = style_or_run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), name)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    set_east_asia(normal)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = HEADING_COLOR
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    set_east_asia(h1)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = HEADING_COLOR
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    set_east_asia(h2)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = HEADING_DARK
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    set_east_asia(h3)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.1
        set_east_asia(style)


def add_page_number(paragraph) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def clean_text(value: str) -> str:
    value = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", value)
    value = value.replace("`", "")
    return value.strip()


def add_code_block(doc: Document, code: str) -> None:
    for line in code.splitlines():
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        set_east_asia(run)
    doc.add_paragraph()


def setup_section(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run("AI智学管家 项目开发计划书")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x5B, 0x6B, 0x7A)
    set_east_asia(run)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_para)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    margins = OxmlElement("w:tblCellMar")
    for name, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_cell_margins(table)

    col_count = len(rows[0])
    usable = 6.5
    if col_count <= 2:
        widths = [usable * 0.32, usable * 0.68]
    elif col_count == 3:
        widths = [usable * 0.24, usable * 0.38, usable * 0.38]
    elif col_count == 4:
        widths = [usable * 0.18, usable * 0.32, usable * 0.26, usable * 0.24]
    else:
        widths = [usable / col_count] * col_count

    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.05
            run = para.add_run(clean_text(cell_text))
            run.font.size = Pt(10)
            set_east_asia(run)
            if row_idx == 0:
                run.font.bold = True
                shade_cell(cell, HEADER_FILL)
            if col_idx < len(widths):
                cell.width = Inches(widths[col_idx])

    for idx, width in enumerate(widths):
        for row in table.rows:
            row.cells[idx].width = Inches(width)

    doc.add_paragraph()


def parse_markdown(source: Path) -> list[tuple[str, object]]:
    blocks: list[tuple[str, object]] = []
    lines = source.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []
    table_buffer: list[list[str]] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                blocks.append(("code", "\n".join(code_lines)))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if not re.fullmatch(r":?-{2,}:?", cells[0]) and not table_buffer:
                table_buffer.append(cells)
            elif table_buffer and re.fullmatch(r":?-{2,}:?", cells[0]):
                i += 1
                continue
            elif table_buffer:
                table_buffer.append(cells)
            if i + 1 >= len(lines) or not lines[i + 1].strip().startswith("|"):
                blocks.append(("table", table_buffer))
                table_buffer = []
            i += 1
            continue
        if table_buffer:
            blocks.append(("table", table_buffer))
            table_buffer = []
        if stripped.startswith("### "):
            blocks.append(("h3", clean_text(stripped[4:])))
        elif stripped.startswith("## "):
            blocks.append(("h2", clean_text(stripped[3:])))
        elif stripped.startswith("# "):
            blocks.append(("title", clean_text(stripped[2:])))
        elif stripped.startswith("> "):
            blocks.append(("quote", clean_text(stripped[2:])))
        elif re.match(r"^\d+\.\s", stripped):
            blocks.append(("numbered", clean_text(re.sub(r"^\d+\.\s", "", stripped))))
        elif stripped.startswith("- "):
            blocks.append(("bullet", clean_text(stripped[2:])))
        elif stripped:
            blocks.append(("paragraph", clean_text(stripped)))
        i += 1
    if table_buffer:
        blocks.append(("table", table_buffer))
    return blocks


def build_docx() -> None:
    source_arg = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SOURCE
    output_arg = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_OUTPUT
    source = source_arg if source_arg.is_absolute() else ROOT / source_arg
    output = output_arg if output_arg.is_absolute() else ROOT / output_arg
    doc = Document()
    setup_section(doc)
    setup_styles(doc)

    blocks = parse_markdown(source)
    for kind, payload in blocks:
        if kind == "title":
            para = doc.add_paragraph()
            run = para.add_run(payload)
            run.font.name = "Calibri"
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
            set_east_asia(run)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(8)
            first_title = False
        elif kind == "h1":
            doc.add_heading(payload, level=1)
        elif kind == "h2":
            doc.add_heading(payload, level=2)
        elif kind == "h3":
            doc.add_heading(payload, level=3)
        elif kind == "paragraph":
            if "：" in payload and payload.startswith(("版本", "日期", "状态", "依据")):
                para = doc.add_paragraph()
                run = para.add_run(payload)
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x5B, 0x6B, 0x7A)
                set_east_asia(run)
                continue
            doc.add_paragraph(payload)
        elif kind == "code":
            add_code_block(doc, payload)
        elif kind == "quote":
            para = doc.add_paragraph()
            run = para.add_run(payload)
            run.italic = True
            set_east_asia(run)
            para.paragraph_format.left_indent = Inches(0.25)
        elif kind == "bullet":
            doc.add_paragraph(payload, style="List Bullet")
        elif kind == "numbered":
            doc.add_paragraph(payload, style="List Number")
        elif kind == "table":
            rows = [row for row in payload if row]
            add_table(doc, rows)

    doc.save(output)
    print(f"saved: {output}")


if __name__ == "__main__":
    build_docx()
