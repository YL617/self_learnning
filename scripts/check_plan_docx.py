"""Structural QA for the generated project plan DOCX."""

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "docs" / "项目开发计划书.docx"
SOURCE = ROOT / "docs" / "项目开发计划书.md"

doc = Document(str(DOCX))
full_text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            full_text += "\n" + cell.text

source = SOURCE.read_text(encoding="utf-8")
keys = [
    "AI智学管家 项目开发计划书",
    "破冰问卷",
    "四重质检",
    "艾宾浩斯",
    "用户协议",
    "知识标签",
    "复习调度",
    "防作弊",
    "里程碑计划",
    "数据模型规划",
    "API 规划",
    "合规与风险",
    "cd backend",
    "pnpm dev",
]

failures = []
for key in keys:
    ok = key in source and key in full_text
    print(f"{'PASS' if ok else 'FAIL'} {key}")
    if not ok:
        failures.append(key)

headings = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} headings={headings}")
residue = [marker for marker in ("```", "**", "`") if marker in full_text]
print("residue=" + (",".join(residue) if residue else "none"))

if failures:
    raise SystemExit(1)
